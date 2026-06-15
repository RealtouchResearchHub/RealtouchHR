"""
RealtouchHR - UKVI Compliance Scanner Service

Premium feature: runs structured compliance checks against all visa-holder
employees, scores the company's overall compliance posture, and generates
PDF/DOCX reports.

Quota: 2 scans per billing month across all subscription plans.
Risk scoring:
  90–100%  → Compliant
  75–89%   → Minor gaps
  50–74%   → Moderate risk
  25–49%   → High risk
  0–24%    → Critical risk

Security: No raw UKVI credentials are stored; all checks are against
employee records already in the database.
"""

from __future__ import annotations

import io
import logging
import uuid
from datetime import date, datetime, timezone, timedelta
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Rules catalogue
# ---------------------------------------------------------------------------

COMPLIANCE_RULES: List[Dict[str, Any]] = [
    # Right to Work
    {
        "rule_code": "RTW-001",
        "category": "right_to_work",
        "title": "Right to Work check on file",
        "description": "A valid Right to Work check must be recorded for every employee.",
        "severity": "critical",
        "check_type": "presence",
    },
    {
        "rule_code": "RTW-002",
        "category": "right_to_work",
        "title": "Right to Work document not expired",
        "description": "Time-limited Right to Work documents must not be past their expiry date.",
        "severity": "critical",
        "check_type": "expiry",
    },
    {
        "rule_code": "RTW-003",
        "category": "right_to_work",
        "title": "Right to Work check date recorded",
        "description": "The date the RTW check was performed must be on file.",
        "severity": "medium",
        "check_type": "presence",
    },
    {
        "rule_code": "RTW-004",
        "category": "right_to_work",
        "title": "Right to Work expiring within 60 days",
        "description": "Advance warning — time-limited RTW documents expiring within 60 days require follow-up action.",
        "severity": "warning",
        "check_type": "expiry_warning",
    },
    # Visa & Immigration
    {
        "rule_code": "VIS-001",
        "category": "visa",
        "title": "Visa type recorded for non-settled workers",
        "description": "Employees on a sponsored or time-limited visa must have their visa type on record.",
        "severity": "high",
        "check_type": "presence",
    },
    {
        "rule_code": "VIS-002",
        "category": "visa",
        "title": "Visa not expired",
        "description": "No employee should be working beyond their visa expiry date.",
        "severity": "critical",
        "check_type": "expiry",
    },
    {
        "rule_code": "VIS-003",
        "category": "visa",
        "title": "Visa expiring within 90 days",
        "description": "Visas expiring within 90 days require renewal action.",
        "severity": "warning",
        "check_type": "expiry_warning",
    },
    {
        "rule_code": "VIS-004",
        "category": "visa",
        "title": "Passport number on file",
        "description": "Passport number must be recorded for sponsored employees.",
        "severity": "medium",
        "check_type": "presence",
    },
    # Certificate of Sponsorship
    {
        "rule_code": "COS-001",
        "category": "sponsorship",
        "title": "CoS reference on file for sponsored employees",
        "description": "Skilled Worker and other sponsored visa holders must have a valid Certificate of Sponsorship reference.",
        "severity": "critical",
        "check_type": "presence",
    },
    {
        "rule_code": "COS-002",
        "category": "sponsorship",
        "title": "CoS not expired",
        "description": "The CoS must not be past its expiry date.",
        "severity": "high",
        "check_type": "expiry",
    },
    {
        "rule_code": "COS-003",
        "category": "sponsorship",
        "title": "SOC code on file for Skilled Worker visa",
        "description": "Standard Occupational Classification (SOC) code must be recorded for Skilled Worker visa holders.",
        "severity": "high",
        "check_type": "presence",
    },
    # Reporting obligations
    {
        "rule_code": "RPT-001",
        "category": "reporting",
        "title": "Nationality recorded",
        "description": "Nationality must be recorded for all employees to support reporting obligations.",
        "severity": "medium",
        "check_type": "presence",
    },
    {
        "rule_code": "RPT-002",
        "category": "reporting",
        "title": "No unresolved UKVI alerts older than 30 days",
        "description": "Open UKVI compliance alerts must be actioned within 30 days.",
        "severity": "high",
        "check_type": "alerts",
    },
    # Salary threshold (Skilled Worker)
    {
        "rule_code": "SAL-001",
        "category": "salary",
        "title": "Skilled Worker minimum salary threshold",
        "description": "Skilled Worker visa holders must be paid at least the minimum salary threshold for their SOC code (£26,200 general minimum from April 2024).",
        "severity": "critical",
        "check_type": "salary_threshold",
    },
    # Documents
    {
        "rule_code": "DOC-001",
        "category": "documents",
        "title": "Employee has supporting compliance documents",
        "description": "Sponsored employees should have visa-related documents uploaded.",
        "severity": "low",
        "check_type": "documents",
    },
    # Payroll / job alignment
    {
        "rule_code": "PAY-001",
        "category": "payroll_alignment",
        "title": "CoS salary aligns with payroll salary",
        "description": "Skilled Worker visa holders must be paid at least the CoS salary. Any significant reduction from the CoS salary may trigger a reporting duty.",
        "severity": "critical",
        "check_type": "cos_salary_match",
    },
    {
        "rule_code": "PAY-002",
        "category": "payroll_alignment",
        "title": "CoS work location matches employee work location",
        "description": "The work location recorded on the CoS must match the employee's actual work location. A change may require an SMS report.",
        "severity": "high",
        "check_type": "cos_location_match",
    },
    # Attendance
    {
        "rule_code": "ATT-001",
        "category": "attendance",
        "title": "No unresolved unauthorised absences for sponsored worker",
        "description": "Sponsors must monitor and report prolonged unauthorised absences (10 or more consecutive days) for sponsored workers.",
        "severity": "high",
        "check_type": "unauthorised_absence",
    },
    # Reporting duty triggers
    {
        "rule_code": "RPT-003",
        "category": "reporting",
        "title": "Sponsored worker has start date recorded",
        "description": "Employers must report to UKVI if a sponsored worker fails to start employment on the agreed start date.",
        "severity": "medium",
        "check_type": "presence",
    },
]

SPONSORED_VISA_TYPES = {
    "skilled_worker", "skilled worker", "intra-company transfer", "intra_company_transfer",
    "global talent", "global_talent", "health and care worker", "health_and_care_worker",
    "senior or specialist worker", "graduate trainee", "temporary worker",
    "t2", "t2 general", "tier 2", "tier2", "tier 2 general",
}

# Minimum Skilled Worker salary (April 2024 onwards)
SW_MINIMUM_SALARY_GBP = 26_200


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _parse_date(val: Optional[str]) -> Optional[date]:
    if not val:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(val[:10], fmt).date()
        except (ValueError, TypeError):
            continue
    return None


def _is_sponsored(employee: Dict[str, Any]) -> bool:
    vt = (employee.get("visa_type") or "").lower().strip()
    return vt in SPONSORED_VISA_TYPES or bool(employee.get("cos_reference") or employee.get("cos_number"))


def _risk_label(score: int) -> str:
    if score >= 90:
        return "compliant"
    if score >= 75:
        return "minor_gaps"
    if score >= 50:
        return "moderate_risk"
    if score >= 25:
        return "high_risk"
    return "critical_risk"


# ---------------------------------------------------------------------------
# Core scanner
# ---------------------------------------------------------------------------

class UKVIComplianceService:
    """Stateless compliance scanning service — call via module-level instance."""

    async def get_scan_quota(self, company_id: str, db: Any) -> Dict[str, Any]:
        """Return current billing-month scan quota for the company."""
        from services.payment_service import UKVI_SCAN_QUOTA_PER_MONTH
        today = date.today()
        period_start = today.replace(day=1)
        next_month = (period_start + timedelta(days=32)).replace(day=1)
        period_end = next_month - timedelta(days=1)

        row = await db.ukvi_scan_entitlements.find_one(
            {"company_id": company_id, "billing_period_start": period_start.isoformat()},
            {"_id": 0},
        )
        scans_used = row["scans_used"] if row else 0
        scans_limit = UKVI_SCAN_QUOTA_PER_MONTH

        return {
            "scans_used": scans_used,
            "scans_limit": scans_limit,
            "scans_remaining": max(0, scans_limit - scans_used),
            "period_start": period_start.isoformat(),
            "period_end": period_end.isoformat(),
            "quota_exceeded": scans_used >= scans_limit,
        }

    async def _increment_scan_quota(self, company_id: str, db: Any) -> None:
        today = date.today()
        period_start = today.replace(day=1)
        row = await db.ukvi_scan_entitlements.find_one(
            {"company_id": company_id, "billing_period_start": period_start.isoformat()},
            {"_id": 0},
        )
        if row:
            await db.ukvi_scan_entitlements.update_one(
                {"company_id": company_id, "billing_period_start": period_start.isoformat()},
                {"$set": {"scans_used": row["scans_used"] + 1, "updated_at": datetime.now(timezone.utc).isoformat()}},
            )
        else:
            next_month = (period_start + timedelta(days=32)).replace(day=1)
            period_end = next_month - timedelta(days=1)
            from services.payment_service import UKVI_SCAN_QUOTA_PER_MONTH
            await db.ukvi_scan_entitlements.insert_one({
                "record_id": f"ent_{uuid.uuid4().hex[:12]}",
                "company_id": company_id,
                "billing_period_start": period_start.isoformat(),
                "billing_period_end": period_end.isoformat(),
                "scans_used": 1,
                "scans_limit": UKVI_SCAN_QUOTA_PER_MONTH,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat(),
            })

    # ------------------------------------------------------------------
    # Run a full compliance scan
    # ------------------------------------------------------------------

    async def run_scan(self, company_id: str, initiated_by: str, db: Any) -> Dict[str, Any]:
        """Execute a full UKVI compliance scan for the company."""
        quota = await self.get_scan_quota(company_id, db)
        if quota["quota_exceeded"]:
            raise ValueError(
                f"UKVI scan quota reached: {quota['scans_used']}/{quota['scans_limit']} scans used "
                f"this billing period (resets {quota['period_start'][:7]} month end)."
            )

        scan_id = f"scan_{uuid.uuid4().hex[:16]}"
        now = datetime.now(timezone.utc)

        # Create scan record
        await db.ukvi_compliance_scans.insert_one({
            "scan_id": scan_id,
            "company_id": company_id,
            "initiated_by": initiated_by,
            "status": "running",
            "created_at": now.isoformat(),
        })

        try:
            results, summary = await self._run_checks(company_id, scan_id, db)
            score = self._calculate_score(results)
            risk = _risk_label(score)

            # Persist results
            for r in results:
                await db.ukvi_compliance_scan_results.insert_one(r)

            # Upsert alerts for failed checks (deduplication: update existing open alert)
            for r in results:
                if r["status"] in ("fail",):
                    existing = await db.ukvi_compliance_alerts.find_one({
                        "company_id": company_id,
                        "employee_id": r.get("employee_id"),
                        "rule_code": r.get("rule_code"),
                        "status": {"$in": ["open", "in_progress"]},
                    }, {"_id": 0})
                    if existing:
                        await db.ukvi_compliance_alerts.update_one(
                            {"alert_id": existing["alert_id"]},
                            {"$set": {"last_seen_scan_id": scan_id, "updated_at": datetime.now(timezone.utc).isoformat(), "description": r.get("message", "")}},
                        )
                    else:
                        await db.ukvi_compliance_alerts.insert_one({
                            "alert_id": f"alrt_{uuid.uuid4().hex[:12]}",
                            "company_id": company_id,
                            "employee_id": r.get("employee_id"),
                            "employee_name": r.get("employee_name"),
                            "rule_code": r.get("rule_code"),
                            "category": r.get("category"),
                            "severity": r.get("severity"),
                            "title": next((rule["title"] for rule in COMPLIANCE_RULES if rule["rule_code"] == r.get("rule_code")), r.get("rule_code")),
                            "description": r.get("message", ""),
                            "status": "open",
                            "source_scan_id": scan_id,
                            "last_seen_scan_id": scan_id,
                            "created_at": datetime.now(timezone.utc).isoformat(),
                            "updated_at": datetime.now(timezone.utc).isoformat(),
                        })

            # Update scan record
            await db.ukvi_compliance_scans.update_one(
                {"scan_id": scan_id},
                {"$set": {
                    "status": "completed",
                    "overall_score": score,
                    "risk_level": risk,
                    "total_employees": summary["total_employees"],
                    "flagged_employees": summary["flagged_employees"],
                    "passed_checks": summary["passed"],
                    "failed_checks": summary["failed"],
                    "warning_checks": summary["warnings"],
                    "summary": summary,
                    "completed_at": datetime.now(timezone.utc).isoformat(),
                }},
            )

            await self._increment_scan_quota(company_id, db)

            # Build inline preview from in-memory results (avoids DB round-trip)
            by_employee: dict = {}
            by_category: dict = {}
            for r in results:
                eid = r.get("employee_id", "unknown")
                if eid not in by_employee:
                    by_employee[eid] = {
                        "employee_id": eid,
                        "employee_name": r.get("employee_name", "Unknown"),
                        "checks": [], "fail_count": 0, "warning_count": 0,
                    }
                by_employee[eid]["checks"].append({
                    "rule_code": r.get("rule_code"),
                    "category": r.get("category"),
                    "status": r.get("status"),
                    "severity": r.get("severity"),
                    "message": r.get("message"),
                })
                if r.get("status") == "fail":
                    by_employee[eid]["fail_count"] += 1
                elif r.get("status") == "warning":
                    by_employee[eid]["warning_count"] += 1
                cat = r.get("category", "other")
                if cat not in by_category:
                    by_category[cat] = {"passed": 0, "failed": 0, "warnings": 0}
                if r.get("status") == "pass":
                    by_category[cat]["passed"] += 1
                elif r.get("status") == "fail":
                    by_category[cat]["failed"] += 1
                else:
                    by_category[cat]["warnings"] += 1

            return {
                "scan_id": scan_id,
                "status": "completed",
                "overall_score": score,
                "risk_level": risk,
                "summary": summary,
                "preview": {
                    "scan_id": scan_id,
                    "overall_score": score,
                    "risk_level": risk,
                    "summary": summary,
                    "employees": list(by_employee.values()),
                    "by_category": by_category,
                },
            }

        except Exception as exc:
            await db.ukvi_compliance_scans.update_one(
                {"scan_id": scan_id},
                {"$set": {"status": "failed", "summary": {"error": str(exc)}}},
            )
            raise

    # ------------------------------------------------------------------
    # Internal check runners
    # ------------------------------------------------------------------

    async def _run_checks(
        self, company_id: str, scan_id: str, db: Any
    ) -> Tuple[List[Dict], Dict]:
        employees = await db.employees.find(
            {"company_id": company_id, "status": {"$nin": ["archived", "leaver"]}},
            {"_id": 0},
        ).to_list(1000)

        ukvi_alerts = await db.ukvi_alerts.find(
            {"company_id": company_id, "status": "open"},
            {"_id": 0},
        ).to_list(500)

        old_alerts = [
            a for a in ukvi_alerts
            if self._alert_age_days(a) > 30
        ]

        cos_records: Dict[str, List] = {}
        all_cos = await db.certificates_of_sponsorship.find(
            {"company_id": company_id, "status": "active"},
            {"_id": 0},
        ).to_list(500)
        for c in all_cos:
            cos_records.setdefault(c.get("employee_id", ""), []).append(c)

        rtw_records: Dict[str, List] = {}
        all_rtw = await db.rtw_checks.find(
            {"company_id": company_id},
            {"_id": 0},
        ).to_list(500)
        for r in all_rtw:
            rtw_records.setdefault(r.get("employee_id", ""), []).append(r)

        results: List[Dict] = []
        flagged_employee_ids: set = set()
        today = date.today()

        for emp in employees:
            eid = emp.get("employee_id", "")
            ename = f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip()
            emp_cos = cos_records.get(eid, [])
            emp_rtw = rtw_records.get(eid, [])
            sponsored = _is_sponsored(emp)

            def _result(rule: Dict, status: str, msg: str, detail: Dict = None) -> Dict:
                if status in ("fail", "critical"):
                    flagged_employee_ids.add(eid)
                return {
                    "result_id": f"res_{uuid.uuid4().hex[:12]}",
                    "scan_id": scan_id,
                    "company_id": company_id,
                    "employee_id": eid,
                    "employee_name": ename,
                    "rule_id": rule["rule_code"],
                    "rule_code": rule["rule_code"],
                    "category": rule["category"],
                    "status": status,
                    "severity": rule["severity"],
                    "message": msg,
                    "detail": detail or {},
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }

            for rule in COMPLIANCE_RULES:
                rc = rule["rule_code"]

                # RTW-001: RTW check on file
                if rc == "RTW-001":
                    if emp_rtw or emp.get("right_to_work"):
                        results.append(_result(rule, "pass", "Right to Work check on file."))
                    else:
                        results.append(_result(rule, "fail", "No Right to Work check recorded.", {"employee": ename}))

                # RTW-002: RTW document not expired
                elif rc == "RTW-002":
                    expiry_str = emp.get("visa_expiry") or (emp_rtw[-1].get("expiry_date") if emp_rtw else None)
                    expiry = _parse_date(expiry_str)
                    rtw_type = emp.get("right_to_work", "")
                    if rtw_type in ("british_citizen", "settled_status", "permanent"):
                        results.append(_result(rule, "pass", "Permanent right to work — no expiry applies."))
                    elif expiry and expiry < today:
                        results.append(_result(rule, "fail", f"Right to Work document expired on {expiry_str}.", {"expiry_date": expiry_str}))
                    elif expiry:
                        results.append(_result(rule, "pass", f"Right to Work valid until {expiry_str}."))
                    else:
                        results.append(_result(rule, "warning", "RTW expiry date not recorded for time-limited worker."))

                # RTW-003: Check date recorded
                elif rc == "RTW-003":
                    has_date = any(r.get("verified_at") or r.get("created_at") for r in emp_rtw)
                    if has_date:
                        results.append(_result(rule, "pass", "RTW check date on file."))
                    else:
                        results.append(_result(rule, "warning", "RTW check date not recorded."))

                # RTW-004: Expiring within 60 days
                elif rc == "RTW-004":
                    expiry_str = emp.get("visa_expiry") or (emp_rtw[-1].get("expiry_date") if emp_rtw else None)
                    expiry = _parse_date(expiry_str)
                    rtw_type = emp.get("right_to_work", "")
                    if rtw_type in ("british_citizen", "settled_status", "permanent"):
                        results.append(_result(rule, "pass", "Permanent right to work — no expiry warning needed."))
                    elif expiry and today <= expiry <= today + timedelta(days=60):
                        results.append(_result(rule, "warning", f"RTW expires in {(expiry - today).days} days ({expiry_str})."))
                    elif expiry:
                        results.append(_result(rule, "pass", f"RTW not expiring within 60 days."))
                    else:
                        results.append(_result(rule, "pass", "No time-limited RTW expiry date on file."))

                # VIS-001: Visa type for non-settled workers
                elif rc == "VIS-001":
                    rtw_type = emp.get("right_to_work", "")
                    if rtw_type in ("british_citizen", "settled_status", "eu_settled", "permanent"):
                        results.append(_result(rule, "pass", "Settled/permanent right to work — visa type not required."))
                    elif emp.get("visa_type"):
                        results.append(_result(rule, "pass", f"Visa type recorded: {emp['visa_type']}."))
                    else:
                        results.append(_result(rule, "warning", "Visa type not recorded for this employee."))

                # VIS-002: Visa not expired
                elif rc == "VIS-002":
                    visa_expiry = _parse_date(emp.get("visa_expiry"))
                    if not emp.get("visa_type"):
                        results.append(_result(rule, "pass", "No visa type on record — check not applicable."))
                    elif visa_expiry and visa_expiry < today:
                        results.append(_result(rule, "fail", f"Visa expired on {emp['visa_expiry']}.", {"expiry_date": emp["visa_expiry"]}))
                    elif visa_expiry:
                        results.append(_result(rule, "pass", f"Visa valid until {emp['visa_expiry']}."))
                    else:
                        results.append(_result(rule, "warning", "Visa expiry date not recorded."))

                # VIS-003: Visa expiring within 90 days
                elif rc == "VIS-003":
                    visa_expiry = _parse_date(emp.get("visa_expiry"))
                    if not emp.get("visa_type"):
                        results.append(_result(rule, "pass", "No visa — check not applicable."))
                    elif visa_expiry and today <= visa_expiry <= today + timedelta(days=90):
                        results.append(_result(rule, "warning", f"Visa expires in {(visa_expiry - today).days} days."))
                    else:
                        results.append(_result(rule, "pass", "Visa not expiring within 90 days."))

                # VIS-004: Passport number on file
                elif rc == "VIS-004":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Non-sponsored employee — passport not mandatory."))
                    elif emp.get("passport_number") or any(r.get("document_number") for r in emp_rtw):
                        results.append(_result(rule, "pass", "Passport number on file."))
                    else:
                        results.append(_result(rule, "warning", "Passport number not recorded for sponsored employee."))

                # COS-001: CoS reference
                elif rc == "COS-001":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Non-sponsored — CoS not required."))
                    elif emp.get("cos_reference") or emp.get("cos_number") or emp_cos:
                        results.append(_result(rule, "pass", "CoS reference on file."))
                    else:
                        results.append(_result(rule, "fail", "No Certificate of Sponsorship reference for sponsored employee."))

                # COS-002: CoS not expired
                elif rc == "COS-002":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Non-sponsored — CoS expiry not applicable."))
                    elif emp_cos:
                        expired = [c for c in emp_cos if _parse_date(c.get("expiry_date")) and _parse_date(c.get("expiry_date")) < today]
                        if expired:
                            results.append(_result(rule, "fail", "Active CoS record has expired.", {"expiry": expired[0].get("expiry_date")}))
                        else:
                            results.append(_result(rule, "pass", "CoS is current."))
                    else:
                        results.append(_result(rule, "warning", "No CoS records to validate expiry against."))

                # COS-003: SOC code
                elif rc == "COS-003":
                    vt = (emp.get("visa_type") or "").lower()
                    is_sw = "skilled" in vt or "t2" in vt or "tier 2" in vt
                    if not is_sw:
                        results.append(_result(rule, "pass", "Skilled Worker visa not applicable."))
                    elif emp.get("soc_code") or any(c.get("soc_code") for c in emp_cos):
                        results.append(_result(rule, "pass", "SOC code on file."))
                    else:
                        results.append(_result(rule, "fail", "SOC code not recorded for Skilled Worker visa holder."))

                # RPT-001: Nationality
                elif rc == "RPT-001":
                    if emp.get("nationality"):
                        results.append(_result(rule, "pass", "Nationality on file."))
                    else:
                        results.append(_result(rule, "warning", "Nationality not recorded."))

                # RPT-002: Old unresolved alerts
                elif rc == "RPT-002":
                    emp_old_alerts = [a for a in old_alerts if a.get("employee_id") == eid]
                    if emp_old_alerts:
                        results.append(_result(rule, "fail", f"{len(emp_old_alerts)} unresolved UKVI alert(s) older than 30 days.", {"count": len(emp_old_alerts)}))
                    else:
                        results.append(_result(rule, "pass", "No stale unresolved UKVI alerts."))

                # SAL-001: Skilled Worker minimum salary
                elif rc == "SAL-001":
                    vt = (emp.get("visa_type") or "").lower()
                    is_sw = "skilled" in vt or "t2" in vt or "tier 2" in vt
                    if not is_sw:
                        results.append(_result(rule, "pass", "Skilled Worker salary threshold not applicable."))
                    else:
                        salary = emp.get("salary") or 0
                        if salary >= SW_MINIMUM_SALARY_GBP:
                            results.append(_result(rule, "pass", f"Salary £{salary:,.0f} meets minimum threshold."))
                        elif salary > 0:
                            results.append(_result(rule, "fail", f"Salary £{salary:,.0f} is below the £{SW_MINIMUM_SALARY_GBP:,} minimum for Skilled Worker visa.", {"salary": salary, "minimum": SW_MINIMUM_SALARY_GBP}))
                        else:
                            results.append(_result(rule, "warning", "Salary not recorded — cannot verify Skilled Worker threshold."))

                # DOC-001: Supporting documents
                elif rc == "DOC-001":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Non-sponsored — document check not critical."))
                    else:
                        results.append(_result(rule, "warning", "Verify that visa-related documents are uploaded in the Documents module."))

                # PAY-001: CoS salary vs payroll salary
                elif rc == "PAY-001":
                    vt = (emp.get("visa_type") or "").lower()
                    is_sw = "skilled" in vt or "t2" in vt or "tier 2" in vt
                    if not is_sw:
                        results.append(_result(rule, "pass", "CoS salary alignment not applicable."))
                    else:
                        cos_salary = emp.get("cos_salary") or 0
                        actual_salary = emp.get("salary") or 0
                        if not cos_salary:
                            results.append(_result(rule, "warning", "CoS salary not recorded — cannot verify alignment."))
                        elif not actual_salary:
                            results.append(_result(rule, "warning", "Payroll salary not recorded — cannot verify CoS alignment."))
                        else:
                            cos_s = float(cos_salary)
                            act_s = float(actual_salary)
                            diff_pct = abs(act_s - cos_s) / cos_s * 100 if cos_s > 0 else 0
                            if act_s < cos_s and diff_pct > 5:
                                results.append(_result(rule, "fail", f"Payroll salary £{act_s:,.0f} is {diff_pct:.0f}% below CoS salary £{cos_s:,.0f}.", {"cos_salary": cos_s, "actual_salary": act_s}))
                            elif diff_pct > 20:
                                results.append(_result(rule, "warning", f"Salary differs from CoS by {diff_pct:.0f}% — verify if a reporting duty is triggered."))
                            else:
                                results.append(_result(rule, "pass", f"Salary £{act_s:,.0f} aligns with CoS salary £{cos_s:,.0f}."))

                # PAY-002: Work location vs CoS work location
                elif rc == "PAY-002":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Location alignment not applicable for non-sponsored."))
                    else:
                        cos_loc = (emp.get("cos_work_location") or "").lower().strip()
                        actual_loc = (emp.get("work_location") or "").lower().strip()
                        if not cos_loc:
                            results.append(_result(rule, "pass", "CoS work location not recorded — check not applicable."))
                        elif not actual_loc:
                            results.append(_result(rule, "warning", "Employee work location not recorded."))
                        elif cos_loc != actual_loc:
                            results.append(_result(rule, "fail", f"Work location changed: CoS says '{emp.get('cos_work_location')}', employee record says '{emp.get('work_location')}'."))
                        else:
                            results.append(_result(rule, "pass", f"Work location matches CoS: {emp.get('work_location')}."))

                # ATT-001: Unauthorised absences
                elif rc == "ATT-001":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Absence monitoring for sponsored workers only."))
                    else:
                        results.append(_result(rule, "warning", "Check absence records for any unauthorised absences exceeding 10 consecutive days."))

                # RPT-003: Start date recorded for sponsored worker
                elif rc == "RPT-003":
                    if not sponsored:
                        results.append(_result(rule, "pass", "Start date reporting trigger not applicable."))
                    elif emp.get("start_date"):
                        results.append(_result(rule, "pass", f"Start date recorded: {emp['start_date']}."))
                    else:
                        results.append(_result(rule, "warning", "Start date not recorded for sponsored worker — needed to report if worker fails to start."))

        summary = {
            "total_employees": len(employees),
            "flagged_employees": len(flagged_employee_ids),
            "passed": sum(1 for r in results if r["status"] == "pass"),
            "failed": sum(1 for r in results if r["status"] == "fail"),
            "warnings": sum(1 for r in results if r["status"] == "warning"),
            "rules_run": len(COMPLIANCE_RULES),
        }
        return results, summary

    # ------------------------------------------------------------------
    # Scoring
    # ------------------------------------------------------------------

    def _calculate_score(self, results: List[Dict]) -> int:
        if not results:
            return 100
        weights = {"critical": 10, "high": 5, "medium": 3, "low": 1, "warning": 2}
        total_weight = 0
        penalty = 0
        for r in results:
            w = weights.get(r.get("severity", "low"), 1)
            total_weight += w
            if r["status"] == "fail":
                penalty += w
            elif r["status"] == "warning":
                penalty += w * 0.3
        if total_weight == 0:
            return 100
        score = max(0, int(100 - (penalty / total_weight) * 100))
        return score

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _alert_age_days(self, alert: Dict) -> int:
        created = alert.get("created_at")
        if not created:
            return 0
        try:
            if isinstance(created, str):
                dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
            else:
                dt = created
            return (datetime.now(timezone.utc) - dt).days
        except Exception:
            return 0

    # ------------------------------------------------------------------
    # Preview (no download charge — returns structured data)
    # ------------------------------------------------------------------

    async def get_scan_preview(self, scan_id: str, company_id: str, db: Any) -> Dict[str, Any]:
        scan = await db.ukvi_compliance_scans.find_one({"scan_id": scan_id, "company_id": company_id}, {"_id": 0})
        if not scan:
            raise ValueError("Scan not found")
        results = await db.ukvi_compliance_scan_results.find(
            {"scan_id": scan_id},
            {"_id": 0},
        ).to_list(2000)

        by_employee: Dict[str, Dict] = {}
        for r in results:
            eid = r.get("employee_id", "unknown")
            if eid not in by_employee:
                by_employee[eid] = {
                    "employee_id": eid,
                    "employee_name": r.get("employee_name", "Unknown"),
                    "checks": [],
                    "fail_count": 0,
                    "warning_count": 0,
                }
            by_employee[eid]["checks"].append({
                "rule_code": r.get("rule_code"),
                "category": r.get("category"),
                "status": r.get("status"),
                "severity": r.get("severity"),
                "message": r.get("message"),
            })
            if r.get("status") == "fail":
                by_employee[eid]["fail_count"] += 1
            elif r.get("status") == "warning":
                by_employee[eid]["warning_count"] += 1

        by_category: Dict[str, Dict] = {}
        for r in results:
            cat = r.get("category", "other")
            if cat not in by_category:
                by_category[cat] = {"passed": 0, "failed": 0, "warnings": 0}
            if r["status"] == "pass":
                by_category[cat]["passed"] += 1
            elif r["status"] == "fail":
                by_category[cat]["failed"] += 1
            elif r["status"] == "warning":
                by_category[cat]["warnings"] += 1

        return {
            "scan_id": scan_id,
            "status": scan.get("status"),
            "overall_score": scan.get("overall_score"),
            "risk_level": scan.get("risk_level"),
            "created_at": scan.get("created_at"),
            "completed_at": scan.get("completed_at"),
            "summary": scan.get("summary", {}),
            "by_category": by_category,
            "employee_details": sorted(by_employee.values(), key=lambda x: x["fail_count"], reverse=True),
        }

    # ------------------------------------------------------------------
    # PDF export
    # ------------------------------------------------------------------

    async def generate_pdf_report(self, scan_id: str, company_id: str, company_name: str, db: Any) -> bytes:
        """Generate a PDF compliance report using ReportLab."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER

        preview = await self.get_scan_preview(scan_id, company_id, db)
        score = preview.get("overall_score", 0)
        risk = preview.get("risk_level", "unknown")
        summary = preview.get("summary", {})

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm,
                                topMargin=20 * mm, bottomMargin=20 * mm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=18, spaceAfter=6)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceAfter=4)
        normal = styles["Normal"]

        risk_color = {
            "compliant": colors.HexColor("#16a34a"),
            "minor_gaps": colors.HexColor("#ca8a04"),
            "moderate_risk": colors.HexColor("#ea580c"),
            "high_risk": colors.HexColor("#dc2626"),
            "critical_risk": colors.HexColor("#7f1d1d"),
        }.get(risk, colors.black)

        elements = []
        elements.append(Paragraph("UKVI Compliance Scan Report", title_style))
        elements.append(Paragraph(f"Company: {company_name}", normal))
        elements.append(Paragraph(f"Scan ID: {scan_id}", normal))
        elements.append(Paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y %H:%M UTC')}", normal))
        elements.append(Spacer(1, 6 * mm))

        # Score summary table
        risk_display = risk.replace("_", " ").title()
        summary_data = [
            ["Overall Score", "Risk Level", "Employees Scanned", "Issues Found"],
            [f"{score}%", risk_display,
             str(summary.get("total_employees", 0)),
             str(summary.get("failed", 0))],
        ]
        summary_table = Table(summary_data, colWidths=[40 * mm, 45 * mm, 45 * mm, 40 * mm])
        summary_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7c3aed")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f5f3ff"), colors.white]),
            ("TEXTCOLOR", (1, 1), (1, 1), risk_color),
            ("FONTNAME", (1, 1), (1, 1), "Helvetica-Bold"),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 6 * mm))

        # Employee breakdown
        elements.append(Paragraph("Employee-Level Findings", h2))
        emp_details = preview.get("employee_details", [])
        flagged = [e for e in emp_details if e["fail_count"] > 0 or e["warning_count"] > 0]

        if flagged:
            for emp in flagged[:50]:  # cap at 50 to keep PDF manageable
                elements.append(Paragraph(f"<b>{emp['employee_name']}</b> — {emp['fail_count']} issue(s), {emp['warning_count']} warning(s)", normal))
                for chk in emp["checks"]:
                    if chk["status"] != "pass":
                        icon = "✗" if chk["status"] == "fail" else "⚠"
                        elements.append(Paragraph(f"  {icon} [{chk['rule_code']}] {chk['message']}", normal))
                elements.append(Spacer(1, 2 * mm))
        else:
            elements.append(Paragraph("No issues found — all checks passed.", normal))

        elements.append(Spacer(1, 8 * mm))
        elements.append(Paragraph(
            "This report was generated by RealtouchHR. It is an internal record-keeping aid "
            "and does not constitute legal advice. Sponsor compliance is the employer's legal responsibility.",
            ParagraphStyle("Disclaimer", parent=normal, fontSize=8, textColor=colors.grey),
        ))

        doc.build(elements)
        return buf.getvalue()

    async def generate_docx_report(self, scan_id: str, company_id: str, company_name: str, db: Any) -> bytes:
        """Generate a DOCX compliance report (fallback to simple format if python-docx absent)."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            # Fallback: return the PDF bytes if python-docx is not installed
            return await self.generate_pdf_report(scan_id, company_id, company_name, db)

        preview = await self.get_scan_preview(scan_id, company_id, db)
        score = preview.get("overall_score", 0)
        risk = (preview.get("risk_level") or "unknown").replace("_", " ").title()
        summary = preview.get("summary", {})

        document = Document()
        document.add_heading("UKVI Compliance Scan Report", 0)
        document.add_paragraph(f"Company: {company_name}")
        document.add_paragraph(f"Scan ID: {scan_id}")
        document.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y %H:%M UTC')}")
        document.add_paragraph(f"Overall Score: {score}%  |  Risk Level: {risk}")
        document.add_paragraph(
            f"Employees Scanned: {summary.get('total_employees', 0)}  |  "
            f"Issues Found: {summary.get('failed', 0)}  |  "
            f"Warnings: {summary.get('warnings', 0)}"
        )

        document.add_heading("Employee Findings", level=1)
        for emp in preview.get("employee_details", []):
            if emp["fail_count"] > 0 or emp["warning_count"] > 0:
                document.add_heading(emp["employee_name"], level=2)
                for chk in emp["checks"]:
                    if chk["status"] != "pass":
                        icon = "FAIL" if chk["status"] == "fail" else "WARN"
                        document.add_paragraph(f"[{icon}] [{chk['rule_code']}] {chk['message']}")

        document.add_paragraph(
            "\nThis report was generated by RealtouchHR. It does not constitute legal advice."
        )

        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()


# Module-level singleton
ukvi_compliance_service = UKVIComplianceService()
